import re

from typing import Any, Optional, Literal
from pydantic import BaseModel
from docx import Document

from .base_file_generator import BaseFileGenerator
from src.prompts.docx_file_generation_prompts import docx_content_creation_prompt, docx_planner_prompt


class AddHeaderArgs(BaseModel):
    text: str
    level: int

class AddFormattedTextArgs(BaseModel):
    markdown_text: str

class AddBulletPointsArgs(BaseModel):
    items: list[str]

class AddNumberedListArgs(BaseModel):
    items: list[str]

class AddHeaderStep(BaseModel):
    function: Literal["add_header"]
    args: AddHeaderArgs

class AddFormattedTextStep(BaseModel):
    function: Literal["add_formatted_text"]
    args: AddFormattedTextArgs

class AddBulletPointsStep(BaseModel):
    function: Literal["add_bullet_points"]
    args: AddBulletPointsArgs

class AddNumberedListStep(BaseModel):
    function: Literal["add_numbered_list"]
    args: AddNumberedListArgs

class OutputStructure(BaseModel):
    steps: list[AddHeaderStep | AddFormattedTextStep | AddBulletPointsStep | AddNumberedListStep]



class DocxFileGenerator(BaseFileGenerator):
    def __init__(self, model: str = "gpt-4.1"):
        super().__init__(model)
        self.document: Document = None

    def generate_file(self, input: str, output_path: str) -> str:
        self.document = Document()
        content = self.generate_content(input)
        plan = self.generate_plan(content)
        self.execute_plan(plan)
        self.document.save(output_path)
        return output_path

    def generate_content(self, input: str) -> str:
        """Call LLM to generate raw document content."""
        return self.call_llm(input=docx_content_creation_prompt.format(input_instructions=input)).text

    def generate_plan(self, content: str) -> OutputStructure:
        """Call LLM to convert raw content into a structured execution plan."""      
        return self.call_llm(
            input=docx_planner_prompt.format(markdown_content=content),
            output_format=OutputStructure
        ).parsed

    def execute_plan(self, plan: OutputStructure):
        """Iterate over structured plan and execute each step."""
        for step in plan.steps:
            getattr(self, step.function)(**step.args.model_dump())     

    def add_header(self, text: str, level: int = 1):
        """Add a header to the document."""
        if self.document:
            self.document.add_heading(text, level=level)

    def add_formatted_text(self, markdown_text: str):
        """Add formatted text (bold, italic) to the document."""
        if not self.document:
            return
        p = self.document.add_paragraph()
        # Simple markdown parsing for **bold** and *italic*
        tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*)", markdown_text)
        for t in tokens:
            if t.startswith("**") and t.endswith("**"):
                run = p.add_run(t[2:-2])
                run.bold = True
            elif t.startswith("*") and t.endswith("*"):
                run = p.add_run(t[1:-1])
                run.italic = True
            else:
                p.add_run(t)

    def add_bullet_points(self, items: list[str]):
        """Add a bullet list to the document."""
        if not self.document:
            return
        for item in items:
            self.document.add_paragraph(item, style="ListBullet")

    def add_numbered_list(self, items: list[str]):
        """Add a numbered list to the document."""
        if not self.document:
            return
        for item in items:
            self.document.add_paragraph(item, style="ListNumber")